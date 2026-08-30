"""
NetSecureX - Report Generation Engine
Produces a professional per-scan Word (.docx) security report:
executive summary, methodology, findings table, risk matrix, recommendations.

Requires: python-docx  (pip install python-docx)
"""

import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from models import Scan
from risk_engine import compute_scan_summary

BASE_DIR = os.path.abspath(os.path.dirname(__file__))
REPORTS_DIR = os.path.join(BASE_DIR, "reports")

SEVERITY_COLORS = {
    "Critical": RGBColor(0xC0, 0x00, 0x00),
    "High": RGBColor(0xE0, 0x6C, 0x00),
    "Medium": RGBColor(0xC8, 0x9B, 0x00),
    "Low": RGBColor(0x2E, 0x7D, 0x32),
}


def _shade_cell(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_col_widths(table, widths_inches):
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            row.cells[idx].width = Inches(width)
    for idx, width in enumerate(widths_inches):
        table.columns[idx].width = Inches(width)


def _add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading


def _add_title_page(doc, scan):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("NetSecureX")
    run.bold = True
    run.font.size = Pt(30)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("Intelligent Network Vulnerability Scanner & Security Dashboard")
    run2.italic = True
    run2.font.size = Pt(14)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Security Assessment Report\n").bold = True
    meta.add_run(f"Scan ID: {scan.id}\n")
    meta.add_run(f"Target: {scan.target_range}\n")
    meta.add_run(f"Scan Date: {scan.start_time.strftime('%d %B %Y, %H:%M UTC')}\n")
    if scan.end_time:
        meta.add_run(f"Duration: {scan.duration_seconds:.1f} seconds\n")
    meta.add_run(f"Report Generated: {datetime.utcnow().strftime('%d %B %Y, %H:%M UTC')}")

    doc.add_page_break()


def _add_disclaimer(doc):
    _add_heading(doc, "Authorization & Scope Statement", level=1)
    p = doc.add_paragraph(
        "This assessment was performed only against systems for which explicit "
        "authorization was confirmed by the operator prior to scanning, in "
        "accordance with the tool's built-in authorization gate. This report is "
        "intended solely for internal security review and remediation planning. "
        "Unauthorized scanning of networks is illegal under applicable computer "
        "misuse legislation."
    )
    p.runs[0].italic = True


def _add_executive_summary(doc, scan, summary):
    _add_heading(doc, "1. Executive Summary", level=1)
    doc.add_paragraph(
        f"NetSecureX performed an automated vulnerability assessment against "
        f"{summary['host_count']} host(s) within the target scope "
        f"'{scan.target_range}'. The scan identified {summary['open_port_count']} "
        f"open port(s) and {summary['vuln_count']} associated finding(s). "
        f"The computed network-wide risk score is "
        f"{summary['network_risk_score']} / 10."
    )

    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Severity"
    hdr[1].text = "Count"

    for sev in ["Critical", "High", "Medium", "Low"]:
        row = table.add_row().cells
        row[0].text = sev
        row[1].text = str(summary["severity_counts"].get(sev, 0))
        _shade_cell(row[0], {
            "Critical": "F8CBCB", "High": "FCE0C6",
            "Medium": "FCF0C6", "Low": "D9EAD3",
        }[sev])

    _set_col_widths(table, [3.0, 1.5])
    doc.add_paragraph()


def _add_methodology(doc):
    _add_heading(doc, "2. Methodology", level=1)
    doc.add_paragraph(
        "The assessment followed a five-stage automated pipeline implemented "
        "by the NetSecureX scanning engine:"
    )
    steps = [
        "Host Discovery — TCP-based reachability probing across the target range/CIDR.",
        "Port Scanning — multithreaded TCP connect scanning across common service ports.",
        "Service & Banner Detection — lightweight banner grabbing to fingerprint service and version.",
        "OS Fingerprinting — heuristic classification from open-port and banner signatures.",
        "Vulnerability Matching — correlation of detected service/version/port against a local "
        "offline vulnerability signature database, followed by CVSS-weighted risk scoring.",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")
    doc.add_paragraph()


def _add_risk_methodology(doc):
    _add_heading(doc, "3. Risk Scoring Methodology", level=1)
    doc.add_paragraph(
        "Each host's risk score is computed as the CVSS-weighted sum of its "
        "findings, adjusted for exposure, and normalized by the number of open "
        "ports on that host:"
    )
    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula.add_run(
        "host_risk = Σ (vuln_cvss_score × exposure_weight) / open_port_count"
    )
    run.italic = True
    run.font.size = Pt(12)

    doc.add_paragraph(
        "exposure_weight is 1.5 for conventionally internet-facing ports "
        "(e.g. 80, 443, 21, 22, 25, 8080, 8443) and 1.0 for internal-only "
        "services. The resulting score is clamped to a 0–10 scale and mapped "
        "to a severity band: Critical (9.0–10.0), High (7.0–8.9), "
        "Medium (4.0–6.9), Low (0.0–3.9)."
    )
    doc.add_paragraph()


def _add_findings(doc, summary):
    _add_heading(doc, "4. Detailed Findings", level=1)

    for row in summary["host_rows"]:
        host = row["host"]
        _add_heading(doc, f"Host: {host.ip}" + (f" ({host.hostname})" if host.hostname else ""), level=2)
        p = doc.add_paragraph()
        p.add_run("OS Guess: ").bold = True
        p.add_run(f"{host.os_guess}    ")
        p.add_run("Risk Score: ").bold = True
        run = p.add_run(f"{row['risk_score']} / 10  ({row['risk_category']})")
        run.font.color.rgb = SEVERITY_COLORS.get(row["risk_category"], RGBColor(0, 0, 0))
        run.bold = True

        if not host.ports:
            doc.add_paragraph("No open ports detected.")
            continue

        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, text in enumerate(["Port", "Service", "Version/Banner", "CVE", "Severity"]):
            hdr[i].text = text

        for port in host.ports:
            if port.vulnerabilities:
                for v in port.vulnerabilities:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(port.port_no)
                    row_cells[1].text = port.service
                    row_cells[2].text = port.version or "-"
                    row_cells[3].text = v.cve_id
                    row_cells[4].text = f"{v.severity} ({v.cvss_score})"
                    _shade_cell(row_cells[4], {
                        "Critical": "F8CBCB", "High": "FCE0C6",
                        "Medium": "FCF0C6", "Low": "D9EAD3",
                    }.get(v.severity, "FFFFFF"))
            else:
                row_cells = table.add_row().cells
                row_cells[0].text = str(port.port_no)
                row_cells[1].text = port.service
                row_cells[2].text = port.version or "-"
                row_cells[3].text = "-"
                row_cells[4].text = "None"

        _set_col_widths(table, [0.7, 1.1, 2.3, 1.4, 1.3])
        doc.add_paragraph()


def _add_recommendations(doc, summary):
    _add_heading(doc, "5. Remediation Recommendations", level=1)
    recs = []
    if summary["severity_counts"].get("Critical", 0) > 0:
        recs.append("Patch or isolate hosts with Critical findings immediately; "
                     "treat as priority-one incidents.")
    if summary["severity_counts"].get("High", 0) > 0:
        recs.append("Schedule patching for High-severity findings within the "
                     "current maintenance window.")
    recs.extend([
        "Disable legacy plaintext protocols (Telnet, unencrypted FTP) in favor "
        "of SSH/SFTP and TLS-wrapped equivalents.",
        "Restrict administrative and database ports (RDP, VNC, MySQL) to "
        "internal management networks or VPN access only.",
        "Re-run NetSecureX periodically and compare scan history to catch "
        "newly exposed services or regressions.",
    ])
    for r in recs:
        doc.add_paragraph(r, style="List Bullet")
    doc.add_paragraph()


def _add_conclusion(doc, summary):
    _add_heading(doc, "6. Conclusion", level=1)
    doc.add_paragraph(
        f"This assessment identified an overall network risk score of "
        f"{summary['network_risk_score']} / 10 across {summary['host_count']} "
        f"host(s). Addressing the Critical and High severity findings above "
        f"should be prioritized to reduce the organization's attack surface. "
        f"Future scans should be scheduled on a recurring basis to track "
        f"remediation progress and detect new exposures."
    )


def generate_report(scan_id):
    os.makedirs(REPORTS_DIR, exist_ok=True)
    scan = Scan.query.get(scan_id)
    summary = compute_scan_summary(scan_id)

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _add_title_page(doc, scan)
    _add_disclaimer(doc)
    doc.add_paragraph()
    _add_executive_summary(doc, scan, summary)
    _add_methodology(doc)
    _add_risk_methodology(doc)
    _add_findings(doc, summary)
    _add_recommendations(doc, summary)
    _add_conclusion(doc, summary)

    filepath = os.path.join(REPORTS_DIR, f"scan_{scan_id}_report.docx")
    doc.save(filepath)
    return filepath
